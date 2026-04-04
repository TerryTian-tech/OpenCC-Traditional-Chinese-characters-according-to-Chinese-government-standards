import os
import re
import difflib
import tempfile
import shutil
import zipfile

from docx import Document
from opencc import OpenCC
from docx.shared import Pt, RGBColor

from text_converter import detect_encoding, safe_read_file

class DocxTraditionalSimplifiedConverter:
    def __init__(self, log_callback, is_cancelled_callback, config='t2gov', preserve_format=True, convert_footnotes=True):
        """
        初始化转换器
        - 't2gov': 繁体转规范繁体
        - 't2new': 繁体旧字形转新字形，但保留异体字不转换
        - 't2gov_keep_simp': 繁体转规范繁体，但保留文档内原有简体字
        - 't2new_keep_simp': 繁体旧字形转新字形，但保留文档内原有简体字和异体字
        - 't2s': 繁体转简体
        - 's2t': 简体转规范繁体
        """
        self.log_callback = log_callback
        self.is_cancelled_callback = is_cancelled_callback
        self.cc = OpenCC(config)
        self.config = config
        self.preserve_format = preserve_format
        self.convert_footnotes = convert_footnotes

    def log(self, msg):
        if self.log_callback:
            self.log_callback(msg)

    def convert_text(self, text):
        """转换文本内容"""
        if text and isinstance(text, str):
            return self.cc.convert(text)
        return text

    def _convert_footnotes_using_zip_manipulation(self, input_path, output_path):
        """
        通过直接操作docx文件（zip格式）来转换脚注和尾注
        这是最可靠的方法，因为它直接修改XML文件
        """
        try:
            # 检查是否已取消
            if self.is_cancelled_callback and self.is_cancelled_callback():
                return False

            # 创建临时目录
            with tempfile.TemporaryDirectory() as temp_dir:
                # 检查是否已取消
                if self.is_cancelled_callback and self.is_cancelled_callback():
                    return False

                # 解压docx文件
                with zipfile.ZipFile(input_path, 'r') as zip_ref:
                    zip_ref.extractall(temp_dir)

                # 检查是否已取消
                if self.is_cancelled_callback and self.is_cancelled_callback():
                    return False

                # 转换脚注XML文件
                footnotes_path = os.path.join(temp_dir, 'word', 'footnotes.xml')
                if os.path.exists(footnotes_path):
                    self._convert_xml_file(footnotes_path)
                    self.log("已转换脚注内容")
                else:
                    self.log("文档中没有脚注")

                # 检查是否已取消
                if self.is_cancelled_callback and self.is_cancelled_callback():
                    return False

                # 转换尾注XML文件（如果有）
                endnotes_path = os.path.join(temp_dir, 'word', 'endnotes.xml')
                if os.path.exists(endnotes_path):
                    self._convert_xml_file(endnotes_path)
                    self.log("已转换尾注内容")

                # 检查是否已取消
                if self.is_cancelled_callback and self.is_cancelled_callback():
                    return False

                # 重新压缩为docx文件
                with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                    for root, dirs, files in os.walk(temp_dir):
                        for file in files:
                            file_path = os.path.join(root, file)
                            # 计算在zip中的相对路径
                            arcname = os.path.relpath(file_path, temp_dir)
                            zipf.write(file_path, arcname)

                return True

        except Exception as e:
            self.log(f"通过zip操作转换脚注时出错: {e}")
            return False

    def _convert_xml_file(self, xml_path):
        """
        转换XML文件（footnotes.xml / endnotes.xml）中的文本内容。

        采用精准正则匹配 <w:t> 标签内文本的方式，仅替换文字内容，
        不对 XML 做解析-重写操作，从而完全保留原始 XML 结构（命名空间声明、
        属性顺序、XML 声明等），杜绝因 ElementTree 重写导致命名空间前缀
        被篡改（如 w: → ns0:）而使 Word 无法解析脚注/尾注的问题。
        """
        try:
            # 检查是否已取消
            if self.is_cancelled_callback and self.is_cancelled_callback():
                return

            with open(xml_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # 检查是否已取消
            if self.is_cancelled_callback and self.is_cancelled_callback():
                return

            # 精准匹配 <w:t>...</w:t> 标签（含带属性的情况如 <w:t xml:space="preserve">）
            def convert_wt_text(match):
                opening_tag = match.group(1)
                text = match.group(2)
                closing_tag = match.group(3)

                # 只在文本包含中文字符时转换，避免无意义的 OpenCC 调用
                if text and any('\u4e00' <= c <= '\u9fff' for c in text):
                    return opening_tag + self.convert_text(text) + closing_tag
                return match.group(0)

            pattern = r'(<w:t(?:\s[^>]*)?>)(.*?)(</w:t>)'
            converted_content = re.sub(pattern, convert_wt_text, content)

            # 检查是否已取消
            if self.is_cancelled_callback and self.is_cancelled_callback():
                return

            with open(xml_path, 'w', encoding='utf-8') as f:
                f.write(converted_content)

        except Exception as e:
            self.log(f"转换XML文件 {xml_path} 时出错: {e}")

    def convert_document(self, input_path, output_path=None):
        """
        转换整个Word文档，根据设置决定是否保留格式和转换脚注
        """
        # 检查是否已取消
        if self.is_cancelled_callback and self.is_cancelled_callback():
            return None

        if output_path is None:
            filename, ext = os.path.splitext(input_path)
            output_path = f"convert_{filename}{ext}"

        self.log(f"开始转换文档: {input_path}")

        # 首先处理脚注和尾注（通过zip操作）- 根据设置决定是否执行
        if self.convert_footnotes:
            temp_output = output_path + ".temp.docx"
            footnote_success = self._convert_footnotes_using_zip_manipulation(input_path, temp_output)

            # 检查是否已取消
            if self.is_cancelled_callback and self.is_cancelled_callback():
                return None

            if footnote_success:
                # 如果脚注转换成功，使用temp文件继续处理其他内容
                processing_file = temp_output
            else:
                # 如果脚注转换失败，使用原始文件
                self.log("脚注转换失败，将只转换正文内容")
                processing_file = input_path
                temp_output = None
        else:
            self.log("跳过脚注和尾注转换")
            processing_file = input_path
            temp_output = None

        try:
            # 检查是否已取消
            if self.is_cancelled_callback and self.is_cancelled_callback():
                if temp_output and os.path.exists(temp_output):
                    os.remove(temp_output)
                return None

            # 读取文档并转换其他内容
            doc = Document(processing_file)

            # 转换正文段落
            self._convert_paragraphs(doc.paragraphs)

            # 检查是否已取消
            if self.is_cancelled_callback and self.is_cancelled_callback():
                if temp_output and os.path.exists(temp_output):
                    os.remove(temp_output)
                return None

            # 转换表格内容
            self._convert_tables(doc.tables)

            # 检查是否已取消
            if self.is_cancelled_callback and self.is_cancelled_callback():
                if temp_output and os.path.exists(temp_output):
                    os.remove(temp_output)
                return None

            # 转换页眉
            for section in doc.sections:
                self._convert_paragraphs(section.header.paragraphs)
                # 转换页眉中的表格
                if hasattr(section.header, 'tables') and section.header.tables:
                    self._convert_tables(section.header.tables)

                # 检查是否已取消
                if self.is_cancelled_callback and self.is_cancelled_callback():
                    if temp_output and os.path.exists(temp_output):
                        os.remove(temp_output)
                    return None

            # 转换页脚
            for section in doc.sections:
                self._convert_paragraphs(section.footer.paragraphs)
                # 转换页脚中的表格
                if hasattr(section.footer, 'tables') and section.footer.tables:
                    self._convert_tables(section.footer.tables)

                # 检查是否已取消
                if self.is_cancelled_callback and self.is_cancelled_callback():
                    if temp_output and os.path.exists(temp_output):
                        os.remove(temp_output)
                    return None

            # 保存最终文档
            doc.save(output_path)
            self.log(f"转换完成，保存至: {output_path}")

            # 清理临时文件
            if temp_output and os.path.exists(temp_output):
                os.remove(temp_output)

            return output_path

        except Exception as e:
            self.log(f"处理文档时出错: {e}")
            # 如果出错，尝试直接复制文件
            if temp_output and os.path.exists(temp_output):
                shutil.copy2(temp_output, output_path)
                if os.path.exists(temp_output):
                    os.remove(temp_output)
                self.log(f"已保存基本转换的文档: {output_path}")
                return output_path
            else:
                # 最后的手段：直接复制原始文件
                shutil.copy2(input_path, output_path)
                self.log(f"转换失败，已复制原始文件到: {output_path}")
                return output_path

    def _convert_paragraphs(self, paragraphs):
        """转换段落集合"""
        for paragraph in paragraphs:
            # 检查是否已取消
            if self.is_cancelled_callback and self.is_cancelled_callback():
                return
            self._convert_paragraph(paragraph)

    def _convert_paragraph(self, paragraph):
        """
        转换单个段落，根据设置决定是否保留格式。
        preserve_format=True 时，采用「整段转换 + 按位置切回」策略：
          先拼接段落内所有 run 的文本，作为完整上下文交给 OpenCC 转换；
          再根据字符位置索引将转换结果精准分配回各个 run，从而既保留
          完整的上下文语义，又完全不触碰 run 的格式属性。
        preserve_format=False 时，同样逐 run 处理以保护脚注引用等结构元素，
          将全部文本集中到第一个有文字的 run 中（丢失 run 级格式），
          跳过空文本 run 以避免 clear_content() 销毁脚注引用。
        """
        if not paragraph.text.strip():
            return

        if self.preserve_format:
            if paragraph.runs:
                self._convert_paragraph_with_context(paragraph)
            else:
                # 极少见的无 run 情况（如段落仅含超链接等非 run 子元素）
                paragraph.text = self.convert_text(paragraph.text)
        else:
            if paragraph.text.strip():
                # 检查是否已取消
                if self.is_cancelled_callback and self.is_cancelled_callback():
                    return

                self._convert_paragraph_simple(paragraph)

    def _convert_paragraph_simple(self, paragraph):
        """
        不保留格式的段落转换。

        与 paragraph.text = ... 不同，此方法在 run 级别操作以保护脚注引用等结构元素：
        - 将段落的完整文本交给 OpenCC 转换（保证上下文语义）
        - 把全部转换后文本写入第一个有文字内容的 run
        - 清空其余有文字的 run 的文本（不触碰空文本 run）
        - 空文本 run（脚注引用 <w:footnoteReference>、制表符 <w:tab/>、
          图片 <w:drawing/> 等）完全不被修改，结构元素得以保留

        效果：run 级格式（粗体、斜体、字号等）丢失，但脚注/尾注引用完好。
        """
        runs = paragraph.runs
        if not runs:
            # 极少见的无 run 情况，回退到整段替换
            paragraph.text = self.convert_text(paragraph.text)
            return

        # 检查是否已取消
        if self.is_cancelled_callback and self.is_cancelled_callback():
            return

        # 拼接所有 run 文本
        texts = [run.text or '' for run in runs]
        full_text = ''.join(texts)

        # 如果整个段落没有中文，跳过转换
        if not any('\u4e00' <= c <= '\u9fff' for c in full_text):
            return

        # 整段文本交给 OpenCC 转换（保持上下文语义）
        converted_full = self.convert_text(full_text)

        # 检查是否已取消
        if self.is_cancelled_callback and self.is_cancelled_callback():
            return

        # 如果转换前后文本相同，无需修改任何 run
        if converted_full == full_text:
            return

        # 找到第一个有实际文本内容的 run
        first_text_run_idx = None
        for i, t in enumerate(texts):
            if t.strip():
                first_text_run_idx = i
                break

        if first_text_run_idx is None:
            # 所有 run 的文本都为空（只有结构元素），无需修改
            return

        # 将全部转换后的文本集中到第一个有文字的 run（格式以该 run 为准，其余 run 格式丢失）
        runs[first_text_run_idx].text = converted_full

        # 清空其余有文字的 run；跳过空文本 run（保护脚注引用、图片等结构元素）
        for i, run in enumerate(runs):
            if i != first_text_run_idx and texts[i]:
                run.text = ''

    def _convert_paragraph_with_context(self, paragraph):
        """
        整段转换 + 按位置切回：
        1. 拼接段落内所有 run 文本，得到完整段落文本
        2. 将完整文本交给 OpenCC 转换（拥有完整上下文）
        3. 由于中文繁简转换几乎都是 1:1 字符映射，转换前后长度一致，
           可按字符位置索引将转换结果精准分配回每个 run
        4. 仅修改 run.text，完全不触碰 run 的格式属性（rPr）
        """
        runs = paragraph.runs
        if not runs:
            return

        # 检查是否已取消
        if self.is_cancelled_callback and self.is_cancelled_callback():
            return

        texts = [run.text or '' for run in runs]
        full_text = ''.join(texts)

        # 如果整个段落没有中文，跳过转换
        if not any('\u4e00' <= c <= '\u9fff' for c in full_text):
            return

        # 1) 计算每个 run 在完整文本中的起止位置（字符索引）
        positions = []
        offset = 0
        for t in texts:
            positions.append((offset, offset + len(t)))
            offset += len(t)

        # 2) 整段文本交给 OpenCC 转换（拥有完整上下文）
        converted_full = self.convert_text(full_text)

        # 检查是否已取消
        if self.is_cancelled_callback and self.is_cancelled_callback():
            return

        # 3) 按位置切分写回各 run（格式完全不动）
        if len(converted_full) == len(full_text):
            # 正常情况：长度一致，直接按位置切割
            for i, run in enumerate(runs):
                start, end = positions[i]
                # 关键：跳过空文本的 run，避免 clear_content() 销毁
                # 脚注引用、制表符、图片等特殊 run 的 text 为空但包含重要子元素
                if start == end:
                    continue
                run.text = converted_full[start:end]
        else:
            # 降级：长度不一致（极其罕见，如某些特殊字符映射）
            # 使用 difflib 逐字符对齐来兜底
            self._convert_paragraph_fallback(runs, texts, positions,
                                             converted_full, full_text)

    def _convert_paragraph_fallback(self, runs, texts, positions, converted_full, full_text):
        """
        降级方案：当转换前后文本长度不一致时，使用 difflib 构建原文字符到
        转换后字符的映射表，按映射关系为每个 run 收集对应的转换后文本。
        """
        # 构建字符级别的差异对齐
        sm = difflib.SequenceMatcher(None, list(full_text), list(converted_full))
        char_map = {}  # original_index -> converted_index

        for op, i1, i2, j1, j2 in sm.get_opcodes():
            if op == 'equal':
                # 等长相同：原位置直接映射到转换后位置
                for k in range(i2 - i1):
                    char_map[i1 + k] = j1 + k
            elif op == 'replace':
                # 替换操作：按最短长度对齐，多出的字符归到最后一个映射位置
                min_len = min(i2 - i1, j2 - j1)
                for k in range(min_len):
                    char_map[i1 + k] = j1 + k
                if (i2 - i1) > min_len:
                    # 原文多出的字符，映射到转换后最后一个对应位置
                    for k in range(min_len, i2 - i1):
                        char_map[i1 + k] = j1 + min_len - 1
                elif (j2 - j1) > min_len:
                    # 转换后多出的字符，追加到最后一个原文对应位置
                    for k in range(min_len, j2 - j1):
                        char_map[i1 + min_len - 1] = j1 + k
            elif op == 'delete':
                # 原文有字符被删除：这些字符映射到前一个有效位置
                for k in range(i1, i2):
                    ref = i1 - 1 if i1 > 0 else 0
                    char_map[k] = char_map.get(ref, j1 if j1 < len(converted_full) else 0)
            elif op == 'insert':
                # 转换后插入了新字符：分配给前一个原文位置
                if i1 > 0 and (i1 - 1) in char_map:
                    char_map[i1 - 1] = j2 - 1

        # 根据映射表为每个 run 收集对应的转换后字符
        for i, run in enumerate(runs):
            start, end = positions[i]
            # 关键：跳过空文本的 run，避免 clear_content() 销毁脚注引用等子元素
            if start == end:
                continue
            converted_chars = []
            for pos in range(start, end):
                if pos in char_map:
                    cpos = char_map[pos]
                    if cpos < len(converted_full):
                        converted_chars.append(converted_full[cpos])
            run.text = ''.join(converted_chars)
    def _convert_tables(self, tables):
        """转换表格内容"""
        for table in tables:
            # 检查是否已取消
            if self.is_cancelled_callback and self.is_cancelled_callback():
                return

            for row in table.rows:
                # 检查是否已取消
                if self.is_cancelled_callback and self.is_cancelled_callback():
                    return

                for cell in row.cells:
                    # 转换单元格中的段落
                    self._convert_paragraphs(cell.paragraphs)

                    # 递归处理嵌套表格
                    for nested_table in cell.tables:
                        self._convert_tables([nested_table])


def convert_docx_file(input_path, output_folder, conversion_type, preserve_format=True, convert_footnotes=True,
                      log_callback=None, is_cancelled_callback=None):
    """
    将Word文档转换为简体/繁体
    :param input_path: 输入文件路径
    :param output_folder: 输出文件夹路径
    :param conversion_type: 转换类型
    :param preserve_format: 是否保留格式
    :param convert_footnotes: 是否转换脚注
    :param log_callback: 日志回调函数
    :param is_cancelled_callback: 取消检查回调函数
    :return: 转换后的文件路径
    """
    def log(msg):
        if log_callback:
            log_callback(msg)

    try:
        # 检查是否已取消
        if is_cancelled_callback and is_cancelled_callback():
            return False

        if not os.path.exists(output_folder):
            os.makedirs(output_folder)
            log(f"创建输出目录: {output_folder}")

        # 检查是否已取消
        if is_cancelled_callback and is_cancelled_callback():
            return False

        if os.path.isfile(input_path) and input_path.lower().endswith('.docx'):
            log(f"正在处理: {os.path.basename(input_path)}")

            # 检查是否已取消
            if is_cancelled_callback and is_cancelled_callback():
                return False

            # 使用新的转换器类
            converter = DocxTraditionalSimplifiedConverter(log_callback, is_cancelled_callback, conversion_type,
                                                            preserve_format, convert_footnotes)
            output_path = os.path.join(output_folder, f"convert_{os.path.basename(input_path)}")
            result = converter.convert_document(input_path, output_path)

            if result:
                log(f"已保存: {result}")
                return result
            else:
                return False

        else:
            log("错误：输入的路径不是有效的.docx文件")
            return False

    except Exception as e:
        log(f"处理 {input_path} 时出错: {str(e)}")
        return False
